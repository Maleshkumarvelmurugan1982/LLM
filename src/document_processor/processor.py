"""
Document processing module for handling various file formats
"""

import os
import uuid
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod
import PyPDF2
from docx import Document
import email
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import re
import asyncio
from pathlib import Path

from ..models.schemas import DocumentType, DocumentChunk, DocumentInfo
from ..utils.logger import get_logger
from ..config import get_settings

logger = get_logger(__name__)

class DocumentProcessor(ABC):
    """Abstract base class for document processors"""
    
    @abstractmethod
    async def process(self, file_path: str, document_id: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Process a document and return chunks"""
        pass
    
    def _create_chunks(self, text: str, document_id: str, chunk_size: int = 1000, 
                      chunk_overlap: int = 200, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Split text into chunks with overlap"""
        if metadata is None:
            metadata = {}
            
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundaries
            if end < len(text):
                # Look for sentence endings
                sentence_ends = ['.', '!', '?', '\n\n']
                for i in range(end, max(start + chunk_size - 100, start), -1):
                    if i < len(text) and text[i] in sentence_ends:
                        end = i + 1
                        break
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunk = DocumentChunk(
                    chunk_id=f"{document_id}_chunk_{chunk_id}",
                    document_id=document_id,
                    text=chunk_text,
                    metadata={**metadata, "chunk_index": chunk_id}
                )
                chunks.append(chunk)
                chunk_id += 1
            
            start = end - chunk_overlap if end - chunk_overlap > start else end
            
        return chunks

class PDFProcessor(DocumentProcessor):
    """PDF document processor"""
    
    async def process(self, file_path: str, document_id: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Process PDF document"""
        try:
            chunks = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            page_metadata = {**metadata, "page_number": page_num + 1}
                            page_chunks = self._create_chunks(
                                text, document_id, 
                                metadata=page_metadata
                            )
                            
                            # Update page number in each chunk
                            for chunk in page_chunks:
                                chunk.page_number = page_num + 1
                                
                            chunks.extend(page_chunks)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        
            logger.info(f"Processed PDF: {len(chunks)} chunks from {len(pdf_reader.pages)} pages")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process PDF {file_path}: {e}")
            raise

class DOCXProcessor(DocumentProcessor):
    """DOCX document processor"""
    
    async def process(self, file_path: str, document_id: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Process DOCX document"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_parts)
            
            chunks = self._create_chunks(full_text, document_id, metadata=metadata)
            
            logger.info(f"Processed DOCX: {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {e}")
            raise

class TXTProcessor(DocumentProcessor):
    """Plain text document processor"""
    
    async def process(self, file_path: str, document_id: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Process text document"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            
            chunks = self._create_chunks(text, document_id, metadata=metadata)
            
            logger.info(f"Processed TXT: {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process TXT {file_path}: {e}")
            raise

class EmailProcessor(DocumentProcessor):
    """Email document processor"""
    
    async def process(self, file_path: str, document_id: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Process email document"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                email_content = file.read()
            
            # Parse email
            msg = email.message_from_string(email_content)
            
            # Extract metadata
            email_metadata = {
                **metadata,
                "subject": msg.get("Subject", ""),
                "from": msg.get("From", ""),
                "to": msg.get("To", ""),
                "date": msg.get("Date", "")
            }
            
            # Extract text content
            text_parts = []
            
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            text_parts.append(payload.decode('utf-8', errors='ignore'))
            else:
                if msg.get_content_type() == "text/plain":
                    payload = msg.get_payload(decode=True)
                    if payload:
                        text_parts.append(payload.decode('utf-8', errors='ignore'))
            
            # Combine header information with body
            header_text = f"Subject: {email_metadata['subject']}\nFrom: {email_metadata['from']}\nTo: {email_metadata['to']}\nDate: {email_metadata['date']}\n\n"
            full_text = header_text + "\n\n".join(text_parts)
            
            chunks = self._create_chunks(full_text, document_id, metadata=email_metadata)
            
            logger.info(f"Processed Email: {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process Email {file_path}: {e}")
            raise

class DocumentProcessorFactory:
    """Factory for creating document processors"""
    
    _processors = {
        DocumentType.PDF: PDFProcessor,
        DocumentType.DOCX: DOCXProcessor,
        DocumentType.TXT: TXTProcessor,
        DocumentType.EMAIL: EmailProcessor,
    }
    
    @classmethod
    def get_processor(cls, document_type: DocumentType) -> DocumentProcessor:
        """Get processor for document type"""
        processor_class = cls._processors.get(document_type)
        if not processor_class:
            raise ValueError(f"Unsupported document type: {document_type}")
        return processor_class()

class DocumentManager:
    """Main document manager"""
    
    def __init__(self):
        self.settings = get_settings()
        self._ensure_directories()
    
    def _ensure_directories(self):
        """Ensure required directories exist"""
        os.makedirs(self.settings.documents_path, exist_ok=True)
        os.makedirs("./logs", exist_ok=True)
    
    async def process_document(self, file_path: str, name: str, document_type: DocumentType, 
                             metadata: Optional[Dict[str, Any]] = None) -> tuple[str, List[DocumentChunk]]:
        """Process a document and return document ID and chunks"""
        
        document_id = str(uuid.uuid4())
        
        if metadata is None:
            metadata = {}
        
        # Add basic metadata
        file_stats = os.stat(file_path)
        metadata.update({
            "document_id": document_id,
            "name": name,
            "document_type": document_type.value,
            "file_size": file_stats.st_size,
            "file_path": file_path
        })
        
        # Get appropriate processor
        processor = DocumentProcessorFactory.get_processor(document_type)
        
        # Process document
        chunks = await processor.process(file_path, document_id, metadata)
        
        logger.info(f"Document {name} processed: {len(chunks)} chunks created")
        
        return document_id, chunks
    
    async def get_document_info(self, file_path: str, name: str, document_type: DocumentType) -> DocumentInfo:
        """Get document information without full processing"""
        
        file_stats = os.stat(file_path)
        
        # Quick page count for PDFs
        page_count = None
        if document_type == DocumentType.PDF:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    page_count = len(pdf_reader.pages)
            except Exception as e:
                logger.warning(f"Could not get page count for PDF: {e}")
        
        return DocumentInfo(
            document_id=str(uuid.uuid4()),
            name=name,
            document_type=document_type,
            file_size=file_stats.st_size,
            page_count=page_count,
            chunk_count=0,  # Will be updated after processing
            metadata={}
        )
